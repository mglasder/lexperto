import glob
import os
from datetime import datetime
from pathlib import Path

from agents import Agent, Runner, trace
from dotenv import load_dotenv
from pydantic import Field, BaseModel

from models.items import (
    AbstrakteErwItem,
    SachverhaltItem,
    BasePromptItem,
    ItemRelevanceDecision,
)
from docx import Document

from models.prompt import PromptBuilder

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

TEST = False


if TEST:
    PROMPTS_FOLDER = "prompts-test"
    TEST_FILENAME_PREFIX = "test_"
else:
    PROMPTS_FOLDER = "prompts"
    TEST_FILENAME_PREFIX = ""


def get_instr(which: str) -> str:
    with open(f"prompts/instructions/{which}", "r", encoding="utf-8") as f:
        instr = f.read()
    return instr


RICHTER_AGENT = Agent(
    name="Richter",
    instructions=get_instr("agent_instruction.txt"),
    model="gpt-4.1-mini",
)


RELEVANCE_DECIDER = Agent(
    name="Relevanzbewerter",
    instructions=get_instr("relevance_decider_instructions.txt"),
    model="gpt-4.1-mini",
    output_type=ItemRelevanceDecision,
)


# TODO: create a detailed prompt with agentic intstructions


def read_file(file_path):
    """Read content from a file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def read_word(file_path):
    """Read content from a Word file."""
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def load_items_to_examine_from(
    folder: str, item_cls: BasePromptItem
) -> list[BasePromptItem]:
    folder_path = Path(folder)
    yaml_files = glob.glob(os.path.join(folder_path, "*.yaml"))

    # sort by numbering in filename
    yaml_files.sort(key=lambda x: int(Path(x).stem.split("_")[1]))

    # exclude template files
    yaml_files = [f for f in yaml_files if "template" not in f]

    items = []

    for yaml_file in yaml_files:
        pp = item_cls.from_yaml(yaml_file)
        items.append(pp)

    return items


def create_sachverhalt(prompt) -> list[str]:
    paragraphs = []

    prompt.add_user(get_instr("sachv_instruction.txt"))

    items = load_items_to_examine_from(
        f"{PROMPTS_FOLDER}/sach/", item_cls=SachverhaltItem
    )

    for item in items:
        prompt.add_user(item.task)
        if item.example:
            prompt.extend(item.examples_as_prompt())
        prompt.add_user("Erstelle jetzt den Paragraphen:")

        result = Runner.run_sync(RICHTER_AGENT, prompt.get())
        paragraphs.append(result.final_output)
        prompt.set_with(result.to_input_list())

    return paragraphs


def create_abstract_considerations(prompt) -> list[str]:
    prompt.add_user(get_instr("aerw_instructions.txt"))

    items = load_items_to_examine_from(
        f"{PROMPTS_FOLDER}/aerw/", item_cls=AbstrakteErwItem
    )
    results = []

    for item in items:
        prompt.add_user("Hier ist die Angabe, was zu prüfen ist:")
        prompt.extend(item.task_as_prompt())

        if not item.mandatory:
            prompt.add_user(
                """
                Dieser Prüfungspunkt ist nicht obligatorisch.
                Prüfe, ob der Prüfungspunkt für den vorliegenden Fall relevant ist und begründe dies.
                Hier sind die Voraussetzungen unter denen der Prüfungspunkt zu prüfen ist:
                """
            )
            prompt.extend(item.requirement_as_prompt())

            prompt.add_user(
                """
                Gehe wie folgt vor:
    
                1. Lese die Angabe zu dem zu prüfenden Punkt und das Beispiel sorgfältig.
                2. Lese die Voraussetzungen, die gegeben sein müssen, damit der Punkt relevant wird.
                3. Lese die Verfügung und die Beschwerde und entscheide, ob die Voraussetzung gegeben sind.
                Du darfst nur Informationen aus der Verfügung und der Beschwerde verwenden.
                DU DARFST KEINE INFORMATIONEN HINZUFÜGEN ODER ERFINDEN.
                5. Überprüfe deine Entscheidung auf Richtigkeit. Ist deine Schlussfolgerung logisch und nachvollziehbar?
                6. Gib an, ob der Prüfungspunkte relevant ist (True) oder nicht (False).
                7. Gib eine Begründung für deine Entscheidung an.
                """
            )
            print(f"Deciding: {item.task}")
            decision = Runner.run_sync(RELEVANCE_DECIDER, prompt.get())

            if not decision.final_output.is_relevant:
                continue
            else:
                print(f"Decision made: {decision.final_output}")
                prompt.add_assistant("Der Prüfungspunkt ist relevant.")
                prompt.add_assistant(decision.final_output.reason)

        prompt.add_user("Hier ist das Beispiel, wie eine Prüfung aussehen könnte:")
        prompt.extend(item.examples_as_prompt())
        prompt.add_user(
            "Erstelle jetzt den spezifischen Absatz für den vorliegenden Fall. Erstelle NUR den Absatz, OHNE sonstigen Erklärungen oder Ergänzungen."
        )

        result = Runner.run_sync(RICHTER_AGENT, prompt.get())
        results.append(result.final_output)
        prompt.set_with(result.to_input_list())

    return results


# TODO: create document saver class
def save_output_word(doc: Document, output_folder: str = "output"):
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc.save(f"{output_folder}/{TEST_FILENAME_PREFIX}{now}_ergebnis.docx")


def main():

    with trace("Amtshilfe Workflow"):
        beschwerde = read_word("input/Beschwerde_Clean_Format.docx")
        verfuegung = read_word("input/Verfuegung_Clean_Format.docx")

        prompt = PromptBuilder()
        prompt.add_user("Hier ist die Verfügung:")
        prompt.add_user(verfuegung)
        prompt.add_user("Hier ist die Beschwerdeschrift:")
        prompt.add_user(beschwerde)

        sachverhalt = create_sachverhalt(prompt)

        prompt.add_user(
            "Hier ist der vollständige Sachverhalt, den du formuliert hast:"
        )

        for sach in sachverhalt:
            prompt.add_assistant(sach)

        erwaegungen = create_abstract_considerations(prompt)

        # TODO: create a document creator class
        doc = Document()
        doc.add_heading("Urteil in Amtshilfeverfahren", level=1)
        doc.add_heading("Sachverhalt:", level=2)
        for sach in sachverhalt:
            doc.add_paragraph(sach, style="List Number")
            doc.add_paragraph("")  # Add a blank line between paragraphs

        doc.add_heading("Das Bundesverwaltungsgericht zieht in Erwägungen:", level=2)
        for erw in erwaegungen:
            doc.add_paragraph(erw, style="List Number")
            doc.add_paragraph("")  # Add a blank line between paragraphs

    save_output_word(doc)

    # convert word document to string for output
    print("\n".join([para.text for para in doc.paragraphs]))


if __name__ == "__main__":
    main()
