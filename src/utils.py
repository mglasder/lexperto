import glob
import os
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
import pdfplumber

from src.models.items import BasePromptItem


def load_prompt(name: str, **subs) -> str:
    """Load a Markdown prompt template and substitute <PLACEHOLDERS>."""
    content = (Path("prompts/production/") / name).read_text()
    for key, val in subs.items():
        content = content.replace(f"<{key}>", str(val))
    return content


def get_instr(which: str) -> str:
    with open(f"prompts/production/instructions/{which}", "r", encoding="utf-8") as f:
        instr = f.read()
    return instr


def read_file(file_path):
    """Read content from a file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def read_word(file_path):
    """Read content from a Word file."""
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def load_items_to_examine_from(
    folder: str, item_cls: BasePromptItem
) -> list[BasePromptItem]:
    folder_path = Path(folder)
    yaml_files = glob.glob(os.path.join(folder_path, "*.yaml"))

    # sort by numbering in filename
    yaml_files.sort(key=lambda x: int(Path(x).stem.split("_")[1]))

    # exclude template files
    yaml_files = [f for f in yaml_files if "template" not in f]

    items = []

    for yaml_file in yaml_files:
        pp = item_cls.from_yaml(yaml_file)
        pp.id = Path(yaml_file).stem  # Set the ID to the filename without extension
        items.append(pp)

    return items


def save_output_word(doc: Document, output_folder: str = "output", test=False):
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    TEST_FILENAME_PREFIX = "test_" if test else ""
    doc.save(f"{output_folder}/{TEST_FILENAME_PREFIX}{now}_ergebnis.docx")


def create_word_document(
    erwaegungen: List[str], sachverhalt: List[str], save=True, test=False
) -> Document:
    # TODO: create a document creator class
    doc = Document()
    doc.add_heading("Urteil in Amtshilfeverfahren", level=1)
    doc.add_heading("Sachverhalt:", level=2)
    for sach in sachverhalt:
        doc.add_paragraph(sach, style="List Number")
        doc.add_paragraph("")  # Add a blank line between paragraphs
    doc.add_heading("Das Bundesverwaltungsgericht zieht in Erwägungen:", level=2)
    for erw in erwaegungen:
        doc.add_paragraph(erw, style="List Number")
        doc.add_paragraph("")  # Add a blank line between paragraphs
    if save:
        save_output_word(doc, test=test)
    return doc


def load_pdf(name: str) -> str:
    """Load a PDF file as text.

    Args:
        name (str): The name of the PDF file without extension.
    Returns:
        str: The text content of the PDF file.
    """

    with pdfplumber.open(name) as pdf:
        full_text = ""
        for page in pdf.pages:
            # Extract text from each page, preserving line breaks
            text = page.extract_text()
            if text:
                full_text += text + "\n\n"

    return full_text


def load_html(name: str) -> str:
    """Load an HTML file as text.

    Args:
        name (str): The name of the HTML file without extension.
    Returns:
        str: The text content of the HTML file.
    """
    with open(f"urteile_html/{name}.html", "r", encoding="utf-8") as f:
        return f.read()
